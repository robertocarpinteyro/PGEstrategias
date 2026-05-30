from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, size=22, bold=True, color=(0,0,0), align=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, bold=bold, size=size, color=color)
    return p

def add_body(doc, text, size=11, color=(50,50,50), align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, space_before=4, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, bold=bold, size=size, color=color, italic=italic)
    return p

def cell_shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

GOLD = (212, 175, 55)
DARK = (15, 15, 30)
WHITE = (255, 255, 255)
GRAY = (80, 80, 90)

# ─── PORTADA ───────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ZENITH MOTORS')
set_font(r, bold=True, size=36, color=GOLD)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after = Pt(8)
r2 = p2.add_run('━' * 34)
set_font(r2, size=13, color=GOLD)

add_heading(doc, 'PROPUESTA DE INVERSION PUBLICITARIA', size=20, color=DARK,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
add_heading(doc, 'Agencia de Seminuevos — Puebla, México', size=13, color=GRAY,
            bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=20)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(10)
p3.paragraph_format.space_after = Pt(2)
r3 = p3.add_run('No estas comprando publicidad.')
set_font(r3, bold=True, size=15, color=DARK, italic=True)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(0)
p4.paragraph_format.space_after = Pt(30)
r4 = p4.add_run('Estas comprando el derecho a dominar tu mercado.')
set_font(r4, bold=True, size=15, color=GOLD, italic=True)

add_body(doc, 'Preparada por: PGE Estrategias', size=11, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=30)
add_body(doc, 'Puebla, Mexico — 2025', size=11, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ─── SECCIÓN 1 ─────────────────────────────────────────────
add_heading(doc, '01. EL PROBLEMA QUE NADIE TE ESTA DICIENDO', size=18, color=DARK,
            space_before=16, space_after=10)

sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(0)
sep.paragraph_format.space_after = Pt(12)
set_font(sep.add_run('─' * 38), size=10, color=GOLD)

add_body(doc, (
    'Hay docenas de agencias de seminuevos en Puebla. Todas tienen inventario. '
    'Todas tienen precios. Todas dicen que son "la mejor opcion".'
), size=12, color=DARK, space_before=6, space_after=6)

add_body(doc, (
    '¿Sabes cual es la unica diferencia entre la agencia que cierra 40 unidades '
    'al mes y la que cierra 8?'
), size=12, color=DARK, space_before=6, space_after=4, bold=True)

p_quote = doc.add_paragraph()
p_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_quote.paragraph_format.space_before = Pt(10)
p_quote.paragraph_format.space_after = Pt(10)
r_q = p_quote.add_run(
    'La que vende 40 controla la narrativa.\n'
    'La que vende 8 espera que el cliente la encuentre.'
)
set_font(r_q, bold=True, size=14, color=GOLD, italic=True)

add_body(doc, (
    'Zenith Motors tiene producto. Tiene precios. Tiene servicio. '
    'Lo que no tiene todavia es el sistema que hace que eso importe.'
), size=12, color=DARK, space_before=6, space_after=6)

add_body(doc, (
    'La pregunta no es si debes invertir en publicidad. La pregunta es si puedes '
    'permitirte seguir invisible mientras tu competencia ocupa el espacio digital '
    'que deberia ser tuyo.'
), size=12, color=DARK, space_before=4, space_after=14)

# ─── SECCIÓN 2 ─────────────────────────────────────────────
add_heading(doc, '02. LA OPORTUNIDAD: LOS DATOS NO MIENTEN', size=18, color=DARK,
            space_before=16, space_after=10)

sep2 = doc.add_paragraph()
sep2.paragraph_format.space_before = Pt(0)
sep2.paragraph_format.space_after = Pt(12)
set_font(sep2.add_run('─' * 38), size=10, color=GOLD)

datos = [
    ('\U0001f697  Mercado en expansion',
     'Mexico es el 4.o mercado automotriz mas grande de America Latina. '
     'El segmento seminuevo crece año con año — y Puebla esta en el centro de ese crecimiento.'),
    ('\U0001f4f1  El comprador decide en linea',
     'Mas del 78% de los compradores de auto investiga en internet ANTES de pisar una agencia. '
     'Si Zenith Motors no aparece en esa busqueda, no existe para ese comprador.'),
    ('\U0001f3af  Tienes 1.7 segundos',
     'El tiempo de atencion promedio en redes sociales es de 1.7 segundos. '
     'Sin video de alta calidad con un gancho en los primeros 3 segundos, tu anuncio es scroll.'),
    ('\U0001f4b8  El costo real de no anunciarse',
     'Con un ticket promedio de $200,000 a $400,000 MXN por unidad, cada venta justifica la '
     'inversion completa. Cada cliente que no llega es una comision que se fue con tu competencia.'),
    ('\U0001f4ca  Video profesional = 3x mas conversiones',
     'Los anuncios con video profesional generan hasta 3 veces mas conversiones que los estaticos '
     'en Meta y Google. Esto no es teoria — es benchmarking de industria.'),
]

for icon_title, desc in datos:
    p_d = doc.add_paragraph()
    p_d.paragraph_format.space_before = Pt(6)
    p_d.paragraph_format.space_after = Pt(2)
    r_t = p_d.add_run(icon_title + '   ')
    set_font(r_t, bold=True, size=11, color=DARK)
    r_desc = p_d.add_run(desc)
    set_font(r_desc, size=11, color=GRAY)

doc.add_paragraph()

add_body(doc, (
    'Antes de gastar un solo peso en pauta, la Opcion A incluye un estudio de mercado real '
    'con datos de tu propio publico. Porque invertir sin datos es apostar. '
    'Invertir con datos es estrategia.'
), size=12, color=DARK, bold=True, space_before=8, space_after=14)

doc.add_page_break()

# ─── OPCIÓN A ──────────────────────────────────────────────
add_heading(doc, 'OPCION A', size=28, color=GOLD, space_before=10, space_after=2)
add_heading(doc, 'PRODUCCION LIVE ACTION + ESTRATEGIA INTEGRAL', size=14, color=DARK,
            space_before=2, space_after=4)

p_inv = doc.add_paragraph()
p_inv.paragraph_format.space_before = Pt(4)
p_inv.paragraph_format.space_after = Pt(4)
set_font(p_inv.add_run('INVERSION TOTAL: $49,000 MXN'), bold=True, size=20, color=DARK)

add_body(doc, 'Pago unico  ·  IVA segun regimen fiscal del cliente', size=10, color=GRAY,
         space_before=2, space_after=8)

p_tag = doc.add_paragraph()
p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tag.paragraph_format.space_before = Pt(8)
p_tag.paragraph_format.space_after = Pt(14)
set_font(p_tag.add_run(
    '"Esto no es marketing. Es el sistema completo que convierte a Zenith Motors\n'
    'en la agencia que Puebla recuerda, busca y elige."'
), bold=True, size=12, color=GOLD, italic=True)

fases_a = [
    ('FASE 1 — INVESTIGACION DE MERCADO Y ESTRATEGIA DIGITAL', [
        'Estudio de mercado cuantitativo via Watti (WhatsApp masivo): datos reales de tu comprador, no suposiciones.',
        'Diseño del instrumento de medicion: preguntas optimizadas para el sector automotriz seminuevo.',
        'Tabulacion, analisis estadistico y reporte ejecutivo de insights con graficos.',
        'Definicion de Buyer Personas basada en datos reales — no en intuicion.',
        'Plan estrategico de embudo completo: Top, Middle y Bottom of Funnel.',
        'Copywriting publicitario persuasivo adaptado por canal y plataforma.',
        'Los ganchos de los primeros 3 segundos y los CTAs se construyen desde la data.',
    ]),
    ('FASE 2 — PREPRODUCCION CINEMATOGRAFICA', [
        'Scouting de locaciones: espacios que refuerzan la identidad de Zenith Motors.',
        'Guion profesional construido sobre los insights del estudio de mercado.',
        'Storyboard / Animatic: cero improvisacion en set. Cada escena planificada.',
        'Casting de talentos que conecten con tu publico objetivo definido.',
        'Revision logistica completa: permisos, transporte, tiempos.',
        'Equipo gama cine: camaras, opticas e iluminacion de nivel cinematografico.',
    ]),
    ('FASE 3 — DIA DE RODAJE', [
        'Dia completo de rodaje en locaciones seleccionadas.',
        '2 operadores de camara: cobertura simultanea, angulos dinamicos, cero baches visuales.',
        '1 operador de drone: las tomas aereas que elevan la percepcion de tu marca.',
        '1 director creativo en set: vision integral que hace coherente cada plano.',
        'Catering para el equipo: porque el rendimiento depende de los detalles.',
    ]),
    ('FASE 4 — POSTPRODUCCION', [
        'Color grading cinematografico: identidad visual que eleva la percepcion de Zenith Motors.',
        'Edicion narrativa: ritmo, musicalizacion y ensamble que mantienen la atencion.',
        'VFX y composicion digital: limpieza visual que separa tu spot de la competencia local.',
        'Integracion de IA como herramienta complementaria sin comprometer la autenticidad.',
    ]),
    ('FASE 5 — DISTRIBUCION: PAID MEDIA EN 4 PLATAFORMAS', [
        'Meta Ads (Facebook & Instagram): campañas de video en Reels y Stories para leads y reconocimiento de marca.',
        'Google Ads (Search & Performance Max): captura de demanda activa de compradores que ya buscan seminuevos.',
        'TikTok Ads: distribucion nativa via Spark Ads o In-Feed para audiencias jovenes.',
        'YouTube Ads: cineminuto en formatos In-stream frente a audiencias In-Market de vehiculos.',
        '1 mes de gestion y optimizacion continua con pruebas A/B.',
        'Reporte de rendimiento: CPA, CTR, ROAS y visualizaciones completadas.',
    ]),
    ('FASE 6 — PRESENCIA FISICA: OOH / ESPECTACULARES', [
        '5 artes estaticos en alta resolucion para espectaculares tradicionales.',
        '5 motion graphics animados (loops de 5-8 s) para pantallas LED en via publica.',
    ]),
]

for fase_title, items in fases_a:
    p_fase = doc.add_paragraph()
    p_fase.paragraph_format.space_before = Pt(10)
    p_fase.paragraph_format.space_after = Pt(4)
    set_font(p_fase.add_run('◉  ' + fase_title), bold=True, size=11, color=DARK)
    for item in items:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_before = Pt(2)
        p_item.paragraph_format.space_after = Pt(2)
        p_item.paragraph_format.left_indent = Cm(1)
        if p_item.runs:
            p_item.runs[0].text = item
            set_font(p_item.runs[0], size=10.5, color=GRAY)

# Entregables A
doc.add_paragraph()
add_heading(doc, 'ENTREGABLES — OPCION A', size=14, color=DARK, space_before=14, space_after=8)

table_a = doc.add_table(rows=1, cols=3)
table_a.style = 'Table Grid'
hdr = table_a.rows[0].cells
hdr[0].text = 'ENTREGABLE'
hdr[1].text = 'DESCRIPCION'
hdr[2].text = 'CANT.'
for cell in hdr:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, bold=True, size=10, color=WHITE)
    cell_shade(cell, '0F0F1E')

rows_a = [
    ('Cineminuto', 'Video horizontal, calidad cinematografica, hasta 1:30 min. Pieza principal.', '1'),
    ('Reels 30 s', 'Highlights del cineminuto adaptados a formato corto para redes.', '4'),
    ('Reels alternativos', 'Hasta 60 s, uno por pilar de comunicacion de la marca.', '4'),
    ('Fotografias de campaña', 'Seleccion con correccion de color profesional.', '50'),
    ('Reporte de mercado', 'Dashboard Watti + estrategia de implementacion.', '2 docs'),
    ('Campañas Paid Media', 'Pautas activas en Meta, Google, TikTok y YouTube.', '1 ecosis.'),
    ('Artes OOH estaticos', 'Archivos finales en alta resolucion para impresion.', '5'),
    ('Artes OOH animados', 'Motion graphics en MP4/MOV para pantallas LED exteriores.', '5'),
]

for i, (ent, desc, cant) in enumerate(rows_a):
    row = table_a.add_row().cells
    row[0].text = ent
    row[1].text = desc
    row[2].text = cant
    fill = 'F5F0E0' if i % 2 == 0 else 'FFFFFF'
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, size=10, color=DARK)
        cell_shade(cell, fill)

# ROI
doc.add_paragraph()
add_heading(doc, '¿POR QUE $49,000 MXN ES UNA CIFRA CONSERVADORA?', size=14, color=DARK,
            space_before=14, space_after=8)
add_body(doc, 'Hagamos el ejercicio que ningun vendedor de publicidad hace contigo.',
         size=12, color=DARK, bold=True)

table_roi = doc.add_table(rows=1, cols=2)
table_roi.style = 'Table Grid'

roi_hdr = table_roi.rows[0].cells
roi_hdr[0].text = 'CONCEPTO'
roi_hdr[1].text = 'CIFRA'
for cell in roi_hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            set_font(r, bold=True, size=10, color=WHITE)
    cell_shade(cell, '0F0F1E')

roi_rows = [
    ('Ticket promedio por seminuevo vendido', '$200,000 — $400,000 MXN'),
    ('Comision promedio de la agencia (10-15%)', '$20,000 — $60,000 MXN por unidad'),
    ('Inversion en esta propuesta (Opcion A)', '$49,000 MXN — una sola vez'),
    ('Para recuperar la inversion necesitas vender...', '1 a 3 unidades adicionales. Ese es el umbral.'),
]

for i, (label, value) in enumerate(roi_rows):
    row = table_roi.add_row().cells
    row[0].text = label
    row[1].text = value
    fills = ['F5F0E0', 'FFFFFF']
    fill = fills[i % 2]
    for j, cell in enumerate(row):
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, size=10.5, color=DARK, bold=(j == 1))
        cell_shade(cell, fill)

add_body(doc, (
    'El video, la estrategia y el estudio de mercado no desaparecen despues de un mes. '
    'Siguen trabajando. Siguen atrayendo clientes. El costo ya no existe. El activo permanece.'
), size=11, color=GRAY, italic=True, space_before=10, space_after=14)

doc.add_page_break()

# ─── OPCIÓN B ──────────────────────────────────────────────
add_heading(doc, 'OPCION B', size=28, color=(80,80,90), space_before=10, space_after=2)
add_heading(doc, 'PRODUCCION CON INTELIGENCIA ARTIFICIAL + ESTRATEGIA INTEGRAL',
            size=13, color=DARK, space_before=2, space_after=4)

p_inv_b = doc.add_paragraph()
p_inv_b.paragraph_format.space_before = Pt(4)
p_inv_b.paragraph_format.space_after = Pt(4)
set_font(p_inv_b.add_run('INVERSION TOTAL: $30,000 MXN'), bold=True, size=20, color=DARK)

add_body(doc, 'Pago unico  ·  Ideal para lanzar y validar con agilidad', size=10, color=GRAY,
         space_before=2, space_after=8)

p_tag_b = doc.add_paragraph()
p_tag_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tag_b.paragraph_format.space_before = Pt(8)
p_tag_b.paragraph_format.space_after = Pt(14)
set_font(p_tag_b.add_run(
    '"La Opcion B no es la version barata.\n'
    'Es la version inteligente para quien quiere resultados ahora y escalar despues."'
), bold=True, size=12, color=GRAY, italic=True)

add_body(doc, (
    'Seamos directos: la Opcion B utiliza herramientas de inteligencia artificial generativa '
    'para producir el material visual. No hay dia de rodaje. No hay drone. No hay equipo en set.'
), size=11, color=DARK, space_before=4, space_after=4)

add_body(doc, (
    'Lo que si incluye — y es exactamente lo mismo que la Opcion A en todo lo que importa:'
), size=11, color=DARK, bold=True, space_before=4, space_after=8)

fases_b = [
    ('FASE 1 — INVESTIGACION DE MERCADO Y ESTRATEGIA DIGITAL', [
        'El punto de partida es identico al de la Opcion A. Primero los datos, luego todo lo demas.',
        'Estudio de mercado cuantitativo via Watti: encuesta masiva por WhatsApp.',
        'Procesamiento, tabulacion y reporte ejecutivo de insights.',
        'Definicion de Buyer Personas basada en datos reales.',
        'Plan estrategico de embudo completo: Top, Middle y Bottom of Funnel.',
        'Copywriting publicitario por canal y plataforma.',
    ]),
    ('FASE 2 — PRODUCCION CON INTELIGENCIA ARTIFICIAL', [
        'Concepto creativo y guion desarrollados con la misma profundidad estrategica que el live action.',
        'Generacion de material visual con herramientas de IA generativa de ultima generacion.',
        'Color grading, edicion, musicalizacion y VFX profesionales sobre el material generado.',
        'Resultados que hace 2 años eran imposibles — hoy son una ventaja competitiva real.',
    ]),
    ('FASE 3 — DISTRIBUCION: PAID MEDIA EN 4 PLATAFORMAS', [
        'Meta Ads, Google Ads, TikTok Ads y YouTube Ads configurados y activos.',
        '1 mes de gestion, pruebas A/B y optimizacion continua.',
        'Reporte de rendimiento: CPA, CTR, ROAS y visualizaciones completadas.',
    ]),
    ('FASE 4 — PRESENCIA FISICA: OOH / ESPECTACULARES', [
        '5 artes estaticos en alta resolucion para espectaculares tradicionales.',
        '5 motion graphics animados (loops de 5-8 s) para pantallas LED en via publica.',
    ]),
]

for fase_title, items in fases_b:
    p_fase = doc.add_paragraph()
    p_fase.paragraph_format.space_before = Pt(10)
    p_fase.paragraph_format.space_after = Pt(4)
    set_font(p_fase.add_run('◉  ' + fase_title), bold=True, size=11, color=DARK)
    for item in items:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_before = Pt(2)
        p_item.paragraph_format.space_after = Pt(2)
        p_item.paragraph_format.left_indent = Cm(1)
        if p_item.runs:
            p_item.runs[0].text = item
            set_font(p_item.runs[0], size=10.5, color=GRAY)

add_heading(doc, '¿CUANDO ELEGIR LA OPCION B?', size=13, color=DARK, space_before=14, space_after=6)

cuando_b = [
    'Cuando necesitas velocidad: los tiempos de produccion son significativamente menores.',
    'Cuando el presupuesto inicial es una restriccion real y quieres demostrar ROI antes de escalar.',
    'Cuando vas a testear mensajes y audiencias antes de invertir en produccion live action.',
    'Cuando el canal prioritario es digital y no necesitas el peso aspiracional del live action.',
]

for c in cuando_b:
    p_c = doc.add_paragraph(style='List Bullet')
    p_c.paragraph_format.space_before = Pt(2)
    p_c.paragraph_format.space_after = Pt(2)
    p_c.paragraph_format.left_indent = Cm(1)
    if p_c.runs:
        p_c.runs[0].text = c
        set_font(p_c.runs[0], size=11, color=DARK)

doc.add_page_break()

# ─── COMPARATIVA ────────────────────────────────────────────
add_heading(doc, 'COMPARATIVA OBJETIVA', size=18, color=DARK, space_before=10, space_after=10)

table_comp = doc.add_table(rows=1, cols=3)
table_comp.style = 'Table Grid'
comp_hdr = table_comp.rows[0].cells
comp_hdr[0].text = 'ELEMENTO'
comp_hdr[1].text = 'OPCION A — $49K'
comp_hdr[2].text = 'OPCION B — $30K'
fills_hdr = ['0F0F1E', 'D4AF37', '505050']
for i, cell in enumerate(comp_hdr):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, bold=True, size=10.5, color=WHITE)
    cell_shade(cell, fills_hdr[i])

comp_rows = [
    ('Produccion de video', 'Live action cinematografico', 'IA generativa + edicion prof.'),
    ('Equipo en set', 'Director + 2 camaras + drone', 'No aplica'),
    ('Tiempo de entrega', 'Mayor (logistica de rodaje)', 'Mas agil'),
    ('Estrategia digital', 'Incluida completa', 'Incluida completa'),
    ('Paid Media (4 plataformas)', 'Incluido', 'Incluido'),
    ('OOH / Espectaculares', '10 piezas (5+5)', '10 piezas (5+5)'),
    ('Impacto aspiracional', 'Maximo (real, tangible)', 'Alto (moderno, eficiente)'),
    ('Inversion', '$49,000 MXN', '$30,000 MXN'),
]

for i, (elem, opt_a, opt_b) in enumerate(comp_rows):
    row = table_comp.add_row().cells
    row[0].text = elem
    row[1].text = opt_a
    row[2].text = opt_b
    fill = 'F0ECD8' if i % 2 == 0 else 'FAFAFA'
    for j, cell in enumerate(row):
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                set_font(r, size=10, color=DARK)
        cell_shade(cell, fill)

doc.add_paragraph()

# ─── CIERRE ─────────────────────────────────────────────────
add_heading(doc, 'LO QUE DECIDE LA DIFERENCIA', size=18, color=DARK, space_before=16, space_after=10)

sep3 = doc.add_paragraph()
sep3.paragraph_format.space_before = Pt(0)
sep3.paragraph_format.space_after = Pt(12)
set_font(sep3.add_run('─' * 38), size=10, color=GOLD)

add_body(doc,
    'Voy a decirte algo que pocas personas en este negocio te diran con honestidad.',
    size=12, color=DARK, bold=True, space_before=6)

add_body(doc, (
    'La mayoria de las agencias de seminuevos va a seguir haciendo exactamente lo mismo el proximo año: '
    'fotos de WhatsApp, posts sin estrategia, precios en imagen de Canva. Eso no es competencia. Eso es ruido.'
), size=12, color=DARK, space_before=6, space_after=6)

add_body(doc, (
    'Zenith Motors tiene la oportunidad de hacer algo diferente ahora. No dentro de seis meses cuando el '
    'mercado este mas saturado. No cuando la competencia ya tomo el espacio digital que tu podrias tener.'
), size=12, color=DARK, space_before=4, space_after=4)

p_now = doc.add_paragraph()
p_now.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_now.paragraph_format.space_before = Pt(12)
p_now.paragraph_format.space_after = Pt(12)
set_font(p_now.add_run('Ahora.'), bold=True, size=24, color=GOLD)

add_body(doc, 'Esta propuesta no es un costo. Es una posicion competitiva.',
         size=13, color=DARK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

add_body(doc, (
    'El video sigue corriendo cuando duermes. La estrategia digital trabaja cuando estas atendiendo clientes. '
    'El estudio de mercado elimina las suposiciones de tu inversion en pauta. '
    'Los espectaculares refuerzan tu marca en la calle.'
), size=12, color=DARK, space_before=4, space_after=6)

add_body(doc,
    'Todo conectado. Todo midiendo. Todo orientado a una sola metrica: mas unidades vendidas.',
    size=12, color=DARK, bold=True, space_before=4, space_after=14)

# Summary
add_heading(doc, 'LAS DOS OPCIONES, UNA VEZ MAS', size=14, color=DARK, space_before=10, space_after=8)

table_final = doc.add_table(rows=1, cols=2)
table_final.style = 'Table Grid'
f_cells = table_final.rows[0].cells

final_data = [
    ('OPCION A\nLive Action + Estrategia Integral\n$49,000 MXN\n\n"La version que domina el mercado."',
     '0F0F1E', WHITE),
    ('OPCION B\nProduccion con IA + Estrategia Integral\n$30,000 MXN\n\n"La version que entra rapido y escala."',
     'D4AF37', DARK),
]

for i, (text, bg, fg) in enumerate(final_data):
    cell = f_cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), bold=True, size=12, color=fg)
    cell_shade(cell, bg)
    cell.paragraphs[0].paragraph_format.space_before = Pt(16)
    cell.paragraphs[0].paragraph_format.space_after = Pt(16)

add_body(doc, (
    'Ambas opciones incluyen estrategia, investigacion de mercado, Paid Media y OOH.\n'
    'La unica diferencia real esta en como se produce el video. Todo lo demas es igual.'
), size=11, color=GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=14)

p_f1 = doc.add_paragraph()
p_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_f1.paragraph_format.space_before = Pt(10)
p_f1.paragraph_format.space_after = Pt(2)
set_font(p_f1.add_run('El momento perfecto para anunciarte ya paso.'),
         bold=True, size=13, color=DARK, italic=True)

p_f2 = doc.add_paragraph()
p_f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_f2.paragraph_format.space_before = Pt(2)
p_f2.paragraph_format.space_after = Pt(16)
set_font(p_f2.add_run('El segundo mejor momento es hoy.'),
         bold=True, size=15, color=GOLD, italic=True)

add_body(doc, '¿Tienes preguntas? ¿Quieres ajustar el alcance? Hablemos.',
         size=12, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
add_body(doc, 'La mejor inversion que puedes hacer hoy es una conversacion.',
         size=12, color=DARK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=20)

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(20)
set_font(p_footer.add_run(
    'PGE Estrategias  ·  pgestrategiasagency@gmail.com  ·  Puebla, Mexico — 2025'
), size=9, color=GRAY)

p_conf = doc.add_paragraph()
p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_conf.paragraph_format.space_before = Pt(2)
set_font(p_conf.add_run(
    '© 2025 — Propuesta confidencial preparada exclusivamente para Zenith Motors.'
), size=9, color=GRAY, italic=True)

out_path = '/home/user/PGEstrategias/Propuesta_Zenith_Motors.docx'
doc.save(out_path)
print('Guardado en:', out_path)
